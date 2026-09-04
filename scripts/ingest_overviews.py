#!/usr/bin/env python3
"""
Curriculum ingestion — Addendum C Section C7.

Reads the school's own Curriculum Overview documents and emits:
  - seed/curriculum.json        rows ready for the curriculum_week table
  - seed/readiness_report.json  what could not be imported, and why

Rules this script enforces, from the spec:
  * Where a .docx and a .pdf claim the same subject/year/semester, the .docx wins.
  * Syllabus references are extracted by pattern. They are never inferred.
    A row with no reference imports as topic-only and is flagged for the HOD.
  * Files labelled with a previous academic year are excluded and reported.
  * Duplicates are reported for a human decision. The script never guesses.
  * Week labels are normalised to the Monday of that school week.

Usage:
    python scripts/ingest_overviews.py "../CURRICULUM OVERVIEWS" --year 2026-27
"""

import argparse, json, re, sys
from pathlib import Path
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    sys.exit("pip install python-docx")

# PyMuPDF reads the 43 overviews that exist only as PDF — no Word version behind
# them (readiness report, "could_not_import"). They carry a real text layer and
# the same WEEK / TOPIC / OBJECTIVES table the .docx overviews do, so a PDF is
# extracted to the same cell grid and run through the identical header_map,
# week_number and split_objectives below. One parser, two container formats —
# never a second set of rules that could drift from the first.
#
# Optional: a deployment with no PyMuPDF still ingests every .docx and reports
# each PDF as unreadable exactly as before, rather than failing to start.
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# Cambridge references: 4Np.01, 9E.02, 4Gt.04, 6Ni.07, 4SLp.02 ...
#
# The strand letters are one or TWO capitals. English Speaking & Listening is
# the two-capital case — 4SLm, 4SLp, 4SLg, 4SLr, 4SLs — and a single [A-Z] here
# silently imported all 34 of CP4 English's Speaking & Listening objectives as
# uncoded. They then could not be counted toward coverage, could not enter a
# work key, and could not be tagged by the evaluation loop, which is where the
# bug surfaced.
#
# Same lesson as the header_map fix: before concluding the school's documents
# are missing a code, check that the parser can express the code's shape.
# Widening to {1,2} adds exactly those 16 references across the corpus and no
# false positives — verified against the full ingest before changing it.
REF = re.compile(r'\b(\d{1,2}[A-Z]{1,2}[a-z]{0,2}\.\d{2})\b')

WORD_NUM = {'one':1,'two':2,'three':3,'four':4,'five':5,'six':6,'seven':7,'eight':8,
            'nine':9,'ten':10,'eleven':11,'twelve':12,'thirteen':13,'fourteen':14,'fifteen':15}

YEAR_GROUPS = ['CP1','CP2','CP3','CP4','CP5','CP6','LS1','LS2','LS3',
               'IGCSE 1','IGCSE 2','AS Level','A Level','AS','EY1','EY2','EY3','Acorns']

SUBJECT_HINTS = {
    'MATH':'Mathematics', 'MATHEMATIC':'Mathematics', 'ENGLISH':'English',
    'SCIENCE':'Science', 'COMPUTING':'Computing', 'ICT':'ICT',
    'GLOBAL PERSPECTIVE':'Global Perspectives', 'G.P':'Global Perspectives', 'GP ':'Global Perspectives',
    'FRENCH':'French', 'ART':'Art and Design', 'MDD':'Music, Dance and Drama',
    'P.E':'Physical Education', 'PHYSICAL EDUCATION':'Physical Education', 'PE ':'Physical Education',
}

# The calendar is the single source of truth (Addendum A Section A9) — overview dates
# are matched to it, never the reverse. It is read from supabase/seed/calendar.json,
# which scripts/load_calendar.mjs also writes to the database.
#
# This used to be a dict written out here, and scripts/seed.mjs held a second one. They
# disagreed from the midterm break onward: the seed counted the break as week 9 and this
# file skipped it, so every week from 9 on meant one date to the importer and another to
# the app. Neither was wrong on its own, which is exactly why nobody noticed.
CALENDAR_PATH = Path(__file__).resolve().parent.parent / 'supabase' / 'seed' / 'calendar.json'

try:
    _calendar = json.loads(CALENDAR_PATH.read_text(encoding='utf-8'))
except FileNotFoundError:
    sys.exit(f"no calendar at {CALENDAR_PATH} — it is the source of week dates")

# Week numbers repeat across semesters, so both are keyed by (semester, week).
WEEK_MONDAYS = {(w['semester'], w['week']): w['commencing'] for w in _calendar['weeks']}
TEACHING_WEEKS = {(w['semester'], w['week']) for w in _calendar['weeks'] if w['type'] == 'teaching'}


def week_monday(semester, week):
    """The Monday of a week, or None where the calendar has no such week.

    An overview that runs to week 16 of a 15 week semester is stating something the
    calendar does not have. That is worth reporting rather than dating silently, so the
    row imports with no date and the readiness report carries it.
    """
    return WEEK_MONDAYS.get((semester or 1, week))


FOLDER_YEAR = {'ACORNS': 'Acorns', 'NURSERY': 'EY1', 'MIDDLE': 'EY2', 'RECEPTION': 'EY3',
               'A LEVEL': 'A Level', 'AS LEVEL': 'AS Level', 'IGCSE 1': 'IGCSE 1', 'IGCSE 2': 'IGCSE 2'}


def find_year(text: str):
    # lookahead, not a word boundary: "cp1OVER VIEW MDD" has no boundary after CP1
    flat = text.upper().replace('_', ' ').replace('CP ', 'CP')
    return next((y for y in YEAR_GROUPS
                 if re.search(rf'(?<![A-Z0-9]){re.escape(y)}(?![0-9])', flat)), None)


def classify(path: Path):
    """
    Infer year group, subject and semester. The filename is tried first; where it
    is silent the containing folder is used, because the folder tree is the one
    piece of structure the school has actually kept consistent.
    """
    name = path.stem.upper().replace('_', ' ')
    year = find_year(name)
    if not year:
        folder = path.parent.name.upper()
        year = next((v for k, v in FOLDER_YEAR.items() if k in folder), None) or find_year(folder)
    subject = next((v for k, v in SUBJECT_HINTS.items() if k in name), None)
    if not subject and year in ('Acorns', 'EY1', 'EY2', 'EY3'):
        subject = 'Early Years (integrated)'   # EY overviews are not split by subject
    if 'SEMESTER 2' in name or ' S2 ' in name or name.startswith('S2'):
        semester = 2
    elif 'ALL YEAR' in name:
        semester = 0                                  # both semesters in one file
    else:
        semester = 1
    prior_year = bool(re.search(r'202[0-5]\s*[-–]\s*202[0-6]', name)) and '2026' not in name.split('202')[-1]
    return year, subject, semester, prior_year


def week_number(cell: str):
    """'Week Three*  7th to 11th' -> 3;  '1 (24th Aug-28th Aug)' -> 1"""
    t = cell.strip().lower()
    m = re.search(r'week\s*(\d{1,2})', t)
    if m:
        return int(m.group(1))
    m = re.search(r'week\s+([a-z]+)', t)
    if m and m.group(1) in WORD_NUM:
        return WORD_NUM[m.group(1)]
    # A bare leading week number in its own column — several PDF overviews label
    # the week cell "1 (24th Aug–28th Aug, 2026)" with no the word "week". Bounded
    # to 1–15 so a stray figure elsewhere cannot be read as a week.
    m = re.match(r'(\d{1,2})\b', t)
    if m and 1 <= int(m.group(1)) <= 15:
        return int(m.group(1))
    return None


def unwrap(lines):
    """
    Rejoin an objective that a PDF broke across visual lines.

    A .docx cell puts a newline between objectives and nowhere else, so its lines
    arrive whole. PyMuPDF has no such thing to go on: it ends a line wherever the
    text ran out of column, and "Research: Gather / information from a range of /
    reliable sources" imported as three objectives. LS3 GP's overview is a PDF and
    its four weeks came in as 125 fragments, which is what the study pack then put
    on its cover, mid-sentence.

    A line continues the one above when that line did not finish a sentence and this
    one opens in lower case. An objective of its own starts with a capital, so a
    .docx's already-whole lines pass through untouched.
    """
    out = []
    for line in lines:
        if out and not re.search(r'[.;:!?]$', out[-1]) and re.match(r'[a-z(]', line):
            out[-1] = f'{out[-1]} {line}'
        else:
            out.append(line)
    return out


def split_objectives(text: str):
    """
    Return [{ref, text}]. A reference is copied verbatim when present.
    When absent the objective still imports — as topic-only, flagged.
    """
    body = re.split(r'\bResources\b', text, flags=re.I)[0]
    lines = unwrap([l.strip(' -•\t') for l in body.split('\n') if l.strip(' -•\t')])
    out = []
    for line in lines:
        if len(line) < 12:
            continue
        m = REF.search(line)
        out.append({'ref': m.group(1) if m else None,
                    'text': REF.sub('', line).strip(' .-') if m else line})
    return out


def docx_tables(path: Path):
    """Every table in a .docx, as a grid of plain-string cells."""
    doc = Document(str(path))
    grids = []
    for table in doc.tables:
        grids.append([[c.text.replace('\xa0', ' ') for c in row.cells] for row in table.rows])
    return grids


def pdf_tables(path: Path):
    """
    Every table in a PDF, as the same grid of plain-string cells docx_tables
    returns. PyMuPDF's cell text keeps the intra-cell newlines that
    split_objectives relies on, so a PDF row reads exactly like a docx row.

    A PDF table extracts with empty spacer columns and, occasionally, a fully
    empty leading row; both are harmless — header_map matches columns by header
    text and skips a row with no week number, so position noise is ignored
    rather than needing to be cleaned out first.
    """
    if fitz is None:
        raise RuntimeError('PyMuPDF not installed — cannot read PDF (pip install pymupdf)')
    grids = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            for tab in page.find_tables().tables:
                grid = [[(c or '') for c in row] for row in tab.extract()]
                if grid:
                    grids.append(grid)
    return grids


def header_map(grid):
    """Find which column holds what. Overviews vary; the header row does not lie."""
    def cells(row):
        return [(c or '').strip().lower() for c in row]
    heads = cells(grid[0])
    # A PDF table can open with one or more fully empty rows before the header;
    # a .docx never does. Skip blank leads, then apply the same header-on-next-row
    # rule the .docx overviews already needed.
    lead = 0
    while lead < len(grid) - 1 and not any(h for h in heads):
        lead += 1
        heads = cells(grid[lead])
    if not any('week' in h for h in heads) and len(grid) > lead + 1:
        heads = cells(grid[lead + 1])
        offset = lead + 2
    else:
        offset = lead + 1
    col = {}
    for i, h in enumerate(heads):
        if 'week' in h and 'week' not in col:  col['week'] = i
        elif 'activit' in h:                   col.setdefault('act', i)
        elif 'resource' in h:                  col.setdefault('res', i)

    # The objectives column is scored, not taken first-come. Several overviews
    # run "UNIT / TOPIC | STRAND / FOCUS | LEARNING OBJECTIVES", and a
    # first-match rule locks onto Unit and reads "Unit 4.1 Historical stories"
    # as the objectives — which is how a document carrying 140 syllabus
    # references imported as topic-only.
    def rank(h):
        if 'objective' in h: return 3
        if 'topic' in h:     return 2
        if 'unit' in h:      return 1
        return 0

    best = max(((rank(h), -i, i) for i, h in enumerate(heads) if rank(h)), default=None)
    if best:
        col['obj'] = best[2]

    return (col, offset) if 'week' in col and 'obj' in col else (None, offset)


def emit_row(wk, obj_text, act_text='', res_cells=None):
    """
    Build one curriculum_week row from the raw cell strings, whatever read them.
    Both readers funnel through here so a PDF week and a .docx week are shaped and
    cleaned identically — objectives split the same way, topic derived the same
    way, resources pulled from a Resources column or, failing that, from the tail
    of the objectives cell the way Science overviews write them.
    """
    objs = split_objectives(obj_text)
    if not objs:
        return None
    topic = objs[0]['text'][:90] if objs[0]['ref'] is None else obj_text.split('\n')[0][:90]
    res = list(res_cells) if res_cells else []
    if not res:
        tail = re.split(r'\bResources\b', obj_text, flags=re.I)
        res = tail[1].split('\n') if len(tail) > 1 else []
    return {
        'week': wk,
        # Filled in once the semester is known (parse_row does not see it), by
        # date_rows() below.
        'week_commencing': None,
        'is_teaching_week': None,
        'topic_label': topic.strip(),
        'objectives': objs,
        'activities': [a.strip(' 0123456789.•') for a in act_text.split('\n') if a.strip()],
        'resources': [x.strip(' -•') for x in res if x.strip(' -•')],
    }


def read_docx_overview(path: Path):
    """Clean-celled grids: map columns by header text (header_map)."""
    rows = []
    for grid in docx_tables(path):
        col, offset = header_map(grid)
        if not col:
            continue
        for r in grid[offset:]:
            cells = [(c or '').replace('\xa0', ' ').strip() for c in r]
            if len(cells) <= max(col.values()):
                continue
            wk = week_number(cells[col['week']])
            if not wk:
                continue
            row = emit_row(
                wk, cells[col['obj']],
                cells[col['act']] if 'act' in col else '',
                cells[col['res']].split('\n') if 'res' in col else None,
            )
            if row:
                rows.append(row)
    return rows


def read_pdf_overview(path: Path):
    """
    PDF grids do not map by header position: PyMuPDF routinely places the "WEEK"
    header cell in a different column from the week values beneath it, and wraps a
    single objectives cell across several physical rows. So columns are found by
    content, and a row is assembled by week marker rather than by grid row —
    tolerant of both the column drift and the fragmentation.

    A week begins at any cell that parses as a week number; every following row
    with no week number appends its objectives-column fragment to that week. The
    rules doing the actual reading — week_number, the REF pattern, split_objectives,
    emit_row — are the same ones the .docx path uses.
    """
    rows = []
    for grid in pdf_tables(path):
        if len(grid) < 2:
            continue
        ncol = max(len(r) for r in grid)
        wk_hits = [0] * ncol
        ref_hits = [0] * ncol
        txt_len = [0] * ncol
        for r in grid:
            for i, c in enumerate(r):
                c = c or ''
                if week_number(c) is not None:
                    wk_hits[i] += 1
                ref_hits[i] += len(REF.findall(c))
                txt_len[i] += len(c)
        if not any(wk_hits):
            continue
        wk_col = max(range(ncol), key=lambda i: wk_hits[i])
        obj_col = max(range(ncol), key=lambda i: (ref_hits[i], txt_len[i]))
        if obj_col == wk_col:
            continue

        def cell(r, i):
            return (r[i] or '') if i < len(r) else ''

        pending = None   # (week, [objective fragments])
        for r in grid:
            wn = week_number(cell(r, wk_col))
            if wn is not None:
                if pending:
                    row = emit_row(pending[0], '\n'.join(pending[1]))
                    if row:
                        rows.append(row)
                pending = (wn, [cell(r, obj_col)])
            elif pending:
                frag = cell(r, obj_col)
                if frag.strip():
                    pending[1].append(frag)
        if pending:
            row = emit_row(pending[0], '\n'.join(pending[1]))
            if row:
                rows.append(row)
    return rows


def read_overview(path: Path):
    reader = read_pdf_overview if path.suffix.lower() == '.pdf' else read_docx_overview
    rows = reader(path)
    notes = [] if rows else ['no week table found — needs a human look']
    return rows, notes


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('root')
    ap.add_argument('--year', default='2026-27')
    ap.add_argument('--out', default='supabase/seed')
    args = ap.parse_args()

    root, out = Path(args.root), Path(args.out)
    out.mkdir(parents=True, exist_ok=True)

    claims = defaultdict(list)          # (year_group, subject, semester) -> [paths]
    excluded, unclassified = [], []

    for path in sorted(root.rglob('*')):
        if path.suffix.lower() not in ('.docx', '.pdf') or path.name.startswith('~$'):
            continue
        if 'delete' in path.stem.lower():
            excluded.append({'file': str(path.relative_to(root)), 'why': 'named DELETE'})
            continue
        yg, subj, sem, prior = classify(path)
        if prior:
            excluded.append({'file': str(path.relative_to(root)), 'why': 'labelled with a previous academic year'})
            continue
        if not yg or not subj:
            unclassified.append({'file': str(path.relative_to(root)),
                                 'why': f'could not read {"year group" if not yg else "subject"} from the filename'})
            continue
        claims[(yg, subj, sem)].append(path)

    # An HOD's conflict decisions, recorded in the app and exported to this file:
    # { "<year_group>|<subject>|<semester>": "<winning file, relative to root>" }.
    # A resolved conflict is not asked again — the named file is used and the others
    # ignored. A resolution naming a file that no longer exists is itself reported,
    # rather than silently falling back to a guess.
    resolutions = {}
    res_path = out / 'conflict_resolutions.json'
    if res_path.exists():
        try:
            resolutions = json.loads(res_path.read_text(encoding='utf-8'))
        except Exception as e:
            print(f"!! ignoring {res_path.name}: {e}")

    weeks, conflicts, failed = [], [], []

    for (yg, subj, sem), paths in sorted(claims.items()):
        docx = [p for p in paths if p.suffix.lower() == '.docx']
        chosen = docx[0] if docx else paths[0]                 # docx wins over pdf
        if len(docx) > 1 or (not docx and len(paths) > 1):
            pick = resolutions.get(f'{yg}|{subj}|{sem}')
            match = next((p for p in paths if str(p.relative_to(root)) == pick), None) if pick else None
            if pick and not match:
                conflicts.append({'year_group': yg, 'subject': subj, 'semester': sem,
                                  'files': [str(p.relative_to(root)) for p in paths],
                                  'needs': f'recorded decision names "{pick}", which is not among these files — re-decide'})
                continue
            if not match:
                conflicts.append({'year_group': yg, 'subject': subj, 'semester': sem,
                                  'files': [str(p.relative_to(root)) for p in paths],
                                  'needs': 'HOD picks which file is current — the importer will not guess'})
                continue
            chosen = match                                     # decision honoured
        try:
            rows, notes = read_overview(chosen)
        except Exception as e:
            failed.append({'year_group': yg, 'subject': subj, 'semester': sem,
                           'file': str(chosen.relative_to(root)), 'why': f'{type(e).__name__}: {e}'})
            continue
        if not rows:
            failed.append({'year_group': yg, 'subject': subj, 'semester': sem,
                           'file': str(chosen.relative_to(root)), 'why': notes[0] if notes else 'no rows'})
            continue
        for r in rows:
            # The Monday, now that the semester is known. A semester of 0 means the file
            # covers both, and its week numbers are read as semester 1's - which is what
            # the school's own "Week 5" means in a whole-year overview.
            week_sem = sem if sem in (1, 2) else 1
            r['week_commencing'] = week_monday(week_sem, r['week'])
            r['is_teaching_week'] = (week_sem, r['week']) in TEACHING_WEEKS
            weeks.append({'academic_year': args.year, 'year_group': yg, 'subject': subj,
                          'semester': sem, 'source_file': str(chosen.relative_to(root)), **r})

    # Count distinct curriculum weeks, not raw emitted rows. A PDF whose week
    # spans a page break is read as two tables and emits the same
    # (year_group, subject, semester, week) twice; an overview running parallel
    # units emits it twice legitimately. Either way the registry holds one week —
    # the seed merges them — so the human-facing readiness numbers must count the
    # week, not the fragment, or "weeks imported" reads far higher than the school
    # has weeks.
    def wkey(w):
        return (w['year_group'], w['subject'], w['semester'], w['week'])
    distinct = {wkey(w): False for w in weeks}
    for w in weeks:
        if any(o['ref'] for o in w['objectives']):
            distinct[wkey(w)] = True
    n_weeks = len(distinct)
    n_coded = sum(distinct.values())

    report = {
        'academic_year': args.year,
        'summary': {
            'weeks_imported': n_weeks,
            'weeks_with_syllabus_refs': n_coded,
            'weeks_topic_only': n_weeks - n_coded,
            'subjects_ready': len({(yg, subj, sem) for (yg, subj, sem, _), c in distinct.items() if c}),
            'source_rows_emitted': len(weeks),
            'duplicate_conflicts': len(conflicts),
            'excluded_files': len(excluded),
            'unreadable': len(failed),
            'unclassified_filenames': len(unclassified),
        },
        'blocking_hod_decision': conflicts,
        'excluded': excluded,
        'could_not_import': failed,
        'unclassified': unclassified,
    }

    (out / 'curriculum.json').write_text(json.dumps(weeks, indent=1), encoding='utf-8')
    (out / 'readiness_report.json').write_text(json.dumps(report, indent=1), encoding='utf-8')

    s = report['summary']
    print(f"imported {s['weeks_imported']} weeks "
          f"({s['weeks_with_syllabus_refs']} with syllabus refs, {s['weeks_topic_only']} topic-only)")
    print(f"{s['duplicate_conflicts']} duplicate conflicts need an HOD decision")
    print(f"{s['excluded_files']} excluded, {s['unreadable']} unreadable, {s['unclassified_filenames']} unclassified")
    print(f"-> {out/'curriculum.json'}\n-> {out/'readiness_report.json'}")


if __name__ == '__main__':
    main()
